# -*- coding: utf-8 -*-

"""
Module implementing MainWindow.
"""

from PyQt5.QtCore import * #pyqtSlot
from PyQt5.QtWidgets import * 
from PyQt5.QtGui import *

#QMainWindow, QApplication

from Ui_maizi_pyqt_test import Ui_MainWindow
from info import Dialog
import webbrowser
import time
import docx

from PyQt5.uic.Compiler.qtproxies import QtGui
'''
可以直接从info.py里直接导入类
'''

#该函数要放在MainWindow类外面
def my_read_docx(filename):
    doc=docx.Document(filename)
    fulltext=[]
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)


class MainWindow(QMainWindow, Ui_MainWindow):
    """
    Class documentation goes here.
    """
    def __init__(self, parent=None):
        """
        Constructor
        
        @param parent reference to the parent widget
        @type QWidget
        """
        super(MainWindow, self).__init__(parent)
        self.setupUi(self)
        self.graphicsView.mousePressEvent=self.my_clicked
        time.sleep(1)
    
    def my_clicked(self, e):
        print('clicked')
        webbrowser.open('www.baidu.com')
    
    @pyqtSlot()
    def on_pushButton_2_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        print('姓名：'+self.lineEdit.text())
        print('电话：'+self.lineEdit_2.text())
        print('邮箱：'+self.lineEdit_3.text())        
        print('公司：'+self.lineEdit_4.text())
        
    
    @pyqtSlot()
    def on_pushButton_3_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_str = self.lineEdit_5.text()
        self.textBrowser.append(my_str)
     
    @pyqtSlot()
    def on_pushButton_4_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        print(self.textBrowser.toPlainText())
        
    @pyqtSlot()
    def on_radioButton_2_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print('you select R2')
        #if self.radioButton_5.isChecked():
        #     print('R5 is checked')
        #elif self.radioButton_4.isChecked():
        #    print('R4 is checked')
        self.radioButton_5.setChecked(True)
    
    @pyqtSlot()
    def on_radioButton_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print('you select R1')
        self.radioButton_4.setChecked(True)
    
    @pyqtSlot()
    def on_radioButton_5_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print('you select R5')
    
    @pyqtSlot()
    def on_radioButton_4_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print('you select R4')
     
    @pyqtSlot(int)   
    def on_dial_valueChanged(self, value):
        """
        Slot documentation goes here.
        
        @param value DESCRIPTION
        @type int
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print(value)
        self.lcdNumber.display(value)
    
    @pyqtSlot(int)
    def on_horizontalSlider_valueChanged(self, value):
        """
        Slot documentation goes here.
        
        @param value DESCRIPTION
        @type int
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print(value)        
        self.label_6.setFont(QFont("Roman times",value,QFont.Bold))
    
    @pyqtSlot(int)
    def on_verticalSlider_valueChanged(self, value):
        """
        Slot documentation goes here.
        
        @param value DESCRIPTION
        @type int
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print(value)
        self.lcdNumber.display(value)
        

    
    @pyqtSlot()
    def on_pushButton_7_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_button =QMessageBox.information(self, 'WARNING','请注意！', )
        

    
    @pyqtSlot()
    def on_pushButton_8_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_button =QMessageBox.question(self, 'Question', '是否保存？')
        
        

    @pyqtSlot()
    def on_pushButton_9_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_button = QMessageBox.warning(self, 'warning', '文字编码方式不同，可能会错误')
        
        

    @pyqtSlot()
    def on_pushButton_10_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_button = QMessageBox.critical(self, '严重警告','不管不行')
        

    
    @pyqtSlot()
    def on_pushButton_11_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_button =QMessageBox.about(self, 'about', '这是一个教程，练手的！')
        

    
    @pyqtSlot()
    def on_pushButton_12_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_button =QMessageBox.aboutQt(self, '关于Qt')
      
        

    
    @pyqtSlot()
    def on_pushButton_13_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_str, ok=QInputDialog.getText(self, '字符串输入框','请在此输入',  QLineEdit.Normal, '请输入信息')
        print(my_str)
        
        

    @pyqtSlot()
    def on_pushButton_14_clicked(self):
        """
        PYQT5使用getInt替代getInteger
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_str, ok=QInputDialog.getInt(self, '输入整数','请输入一个整数', 30, 0, 100)
        print(my_str)
        
        

    
    @pyqtSlot()
    def on_pushButton_15_clicked(self):
        """
        Qstring和Qstringlist不再存在，可以直接使用python list替代
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        #my_list = QStringList()
        my_list=['apple','banana', 'pear' ]
        my_str, ok=QInputDialog.getItem(self, '下拉框', '请选择', my_list)
        print(my_str)
        

    @pyqtSlot()
    def on_pushButton_16_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_info = Dialog()       
        my_info.exec_()
        

    @pyqtSlot()
    def on_pushButton_17_clicked(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        self.label_7.setStyleSheet("image: url(:/pic/3.jpg);")
        
        

    
    @pyqtSlot()
    def on_actionOpen_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        print('OPEN!')

        my_file_path,my_file_extension=QFileDialog.getOpenFileName(self,'打开文件','./')
        #教学视频中加了QtGui，由于已默认导入所有的，故去掉了；文件格式后缀的星号不可少,加了过滤器后传入了后缀名，建议以额外参数接收
#         print(my_file_path)
        if my_file_path[-4:]=='.doc' or my_file_path[-5:]=='.docx': 
            #docx库写的函数           
            full_text = my_read_docx(my_file_path.replace("\/", "\\"))
            self.textBrowser_2.append(full_text)           
# win32com库        
# from win32com import client as wc
#         word =wc.Dispatch('Word.Application')
#         word.Visible =0
#             print(my_file_path)            
#             my_worddoc= word.Documents.Open(my_file_path.replace("\/", "\\" ))  #路径需要替换为反斜杠
#           
#             my_count =my_worddoc.Paragraphs.count
#             for i  in range(my_count):
#                 my_pr =my_worddoc.Paragraphs[i].Range
#                 print(my_pr.text)    #.text即可，不用text（）
#                 self.textBrowser_2.append(my_pr.text)            
#             my_worddoc.Close()#Close的C要大写，否则无法关闭

# docx库
#             doc = docx.Document(my_file_path.replace("\/", "\\"))
# #             print(doc.paragraphs[0].text)#paragraphs,不是paragraph
#             for my_paragraph in doc.paragraphs:
#                 self.textBrowser_2.append(my_paragraph.text)                
        elif my_file_path[-4:]=='.txt':
            #         print(my_file_path)
            with open(my_file_path) as f:
                my_data = f.read()
                self.textBrowser_2.append(my_data)
                
        elif my_file_path.endswith('.xlsx') or my_file_path.endswith('.xls'):
#             print('excel')
#使用Xlrd，使用Xlrd只能读取，不能写入，写入需要用Xlwt
            from xlrd import open_workbook
            wb=open_workbook(my_file_path.replace("\/","\\"))
            for s in wb.sheets():
                for row in range(s.nrows):
                    for col in range(s.ncols):
                        print(s.cell(row,col).value)
                


#             #使用win32库
#             from win32com import client as wc
#             excel =wc.Dispatch('Excel.Application')
#             excel.Visible =0
#             my_excel =excel.Workbooks.Open(my_file_path.replace("\/","\\"))
#             print(my_excel.sheets.Count)
#             my_sheet=my_excel.Sheets('sheet1')
#             print(my_sheet.UsedRange.Rows.Count)
#             print(my_sheet.UsedRange.Columns.Count)
#             for i in range(my_sheet.UsedRange.Rows.Count):
#                 for j in range(my_sheet.UsedRange.Columns.Count):                    
#                     value =my_sheet.Cells(i+1,j+1).Value
#                     print(value)                    
#             my_excel.Close()
#             excel.Quit()
            
                           
        else:
            QMessageBox.Information(self,'提示','不支持的文件类型')
#             print(my_data)
        
     
        
                
        

    
    @pyqtSlot()
    def on_actionQuit_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        sys.exit(0)
    
    @pyqtSlot()
    def on_actionAbout_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
        my_button =QMessageBox.about(self, 'about', '这是一个教程，练手的！')
        
        
        

    
    @pyqtSlot()
    def on_actionIntroduction_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        my_button =QMessageBox.information(self, 'WARNING','请注意！', )
        

    @pyqtSlot()
    def on_actionSave_triggered(self):
        """
        Slot documentation goes here.
        """
        # TODO: not implemented yet
        #raise NotImplementedError
#         print('SAVE!')
        my_data = self.textBrowser_2.toPlainText() 
        my_file_path,my_file_extension=QFileDialog.getSaveFileName(self, "'文件另存为'", "./",) 
        #多种文件名时记得以双引号来引起来，用两个分号分隔       
#         print(my_file)
#         print(my_file_extension)
        if my_file_path[-4:]=='.doc' or my_file_path[-5:]=='.docx': 
            doc =docx.Document()
            doc.add_paragraph(my_data)
            doc.save(my_file_path)
            my_button = QMessageBox.information(self, '提示','文件保存成功！', )
            
        elif  my_file_path[-4:]=='.txt':   
            with open(my_file_path,'a+') as  f:             
                    f.write(my_data)
            my_button = QMessageBox.information(self, '提示','文件保存成功！', ) 
            
#         elif my_file_path
        else:
            my_button = QMessageBox.information(self, '提示','不支持的文件格式', )     
                        
if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)   
    #ui = MainWindow()
    #ui.show()
    #sys.exit(app.exec_())
    splash =QSplashScreen(QPixmap(":/pic/5.jpg"))
    splash.show()
    splash.showMessage('正在加载图片资源……',Qt.AlignBottom,Qt.white)
    time.sleep(1)
    splash.showMessage('正在加载音频资源……',Qt.AlignBottom,Qt.white)
    time.sleep(1)
    splash.showMessage('正在渲染界面……',Qt.AlignBottom,Qt.white)
    time.sleep(1)
    app.processEvents()   
    ui = MainWindow()
    ui.show()
    splash.finish(ui)
    sys.exit(app.exec_())
    
